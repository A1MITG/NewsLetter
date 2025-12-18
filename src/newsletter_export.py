"""
CXO Insurance & BFSI Intelligence Brief - Word Export
Production-Ready Executive Document v1.0
"""

from fastapi.responses import Response
import datetime
import pytz
from .sources import scan_all_sources
from .filter import filter_and_rank_stories
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

IST = pytz.timezone('Asia/Kolkata')


def add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph in Word."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '3B82F6')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_section_header(doc, number, title):
    """Add a numbered section header."""
    para = doc.add_paragraph()
    num_run = para.add_run(f"{number:02d}  ")
    num_run.font.size = Pt(14)
    num_run.font.color.rgb = RGBColor(30, 58, 95)
    num_run.font.bold = True
    title_run = para.add_run(title)
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(30, 58, 95)
    title_run.font.bold = True
    para.space_after = Pt(12)
    return para


def get_word_response():
    now = datetime.datetime.now(IST)
    stories = scan_all_sources()
    top_stories, themes, risks, competitive, tech_signals = filter_and_rank_stories(stories)
    date_str = now.strftime('%A, %d %B %Y')
    time_str = now.strftime('%H:%M') + " IST"
    
    doc = Document()
    
    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # ─────────────────────────────────────────────────────────────
    # HEADER SECTION
    # ─────────────────────────────────────────────────────────────
    
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_run = brand.add_run('CONTINUUM INTELLIGENCE')
    brand_run.font.size = Pt(10)
    brand_run.font.color.rgb = RGBColor(59, 130, 246)
    brand_run.font.bold = True
    brand_run.font.all_caps = True
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('Daily Insurance & BFSI CXO Intelligence Brief')
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(30, 58, 95)
    title_run.font.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run('Strategic Insights for Executive Decision-Makers')
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(100, 116, 139)
    sub_run.font.italic = True
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"{date_str}  •  {time_str}  •  Coverage: Last 24 Hours")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph('─' * 70)
    
    # ─────────────────────────────────────────────────────────────
    # 01. EXECUTIVE SNAPSHOT
    # ─────────────────────────────────────────────────────────────
    
    add_section_header(doc, 1, "Executive Snapshot")
    
    snap_intro = doc.add_paragraph()
    snap_run = snap_intro.add_run("60-Second Read")
    snap_run.font.size = Pt(9)
    snap_run.font.color.rgb = RGBColor(59, 130, 246)
    snap_run.font.bold = True
    snap_run.font.all_caps = True
    
    # Generate executive bullets
    bullets = generate_executive_bullets(top_stories, themes, risks, tech_signals)
    for bullet in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        br = bp.add_run(bullet["text"])
        br.font.size = Pt(11)
        if bullet.get("cxo"):
            bp.add_run(f"  [{bullet['cxo']}]")
    
    doc.add_paragraph()
    
    # ─────────────────────────────────────────────────────────────
    # 02. STRATEGIC THEMES
    # ─────────────────────────────────────────────────────────────
    
    add_section_header(doc, 2, "Strategic Themes Emerging Today")
    
    for theme in themes:
        tp = doc.add_paragraph()
        theme_name = tp.add_run(f"▸ {theme['name']}")
        theme_name.font.size = Pt(12)
        theme_name.font.bold = True
        theme_name.font.color.rgb = RGBColor(30, 58, 95)
        
        insight_p = doc.add_paragraph()
        insight_r = insight_p.add_run(f"   {theme['insight']}")
        insight_r.font.size = Pt(10)
        insight_r.font.color.rgb = RGBColor(100, 116, 139)
        insight_p.space_after = Pt(8)
    
    doc.add_paragraph()
    
    # ─────────────────────────────────────────────────────────────
    # 03. CURATED INTELLIGENCE
    # ─────────────────────────────────────────────────────────────
    
    add_section_header(doc, 3, "Curated Intelligence")
    
    for i, s in enumerate(top_stories[:10], 1):
        # Headline with number
        h_para = doc.add_paragraph()
        h_num = h_para.add_run(f"{i}. ")
        h_num.font.size = Pt(12)
        h_num.font.bold = True
        h_num.font.color.rgb = RGBColor(30, 58, 95)
        h_text = h_para.add_run(s["headline"])
        h_text.font.size = Pt(12)
        h_text.font.bold = True
        
        # Meta tags line
        meta_p = doc.add_paragraph()
        src_r = meta_p.add_run(f"{s.get('source', 'Source')}  •  ")
        src_r.font.size = Pt(9)
        src_r.font.color.rgb = RGBColor(79, 70, 229)
        theme_r = meta_p.add_run(f"{s.get('theme', 'General')[:25]}  •  ")
        theme_r.font.size = Pt(9)
        theme_r.font.color.rgb = RGBColor(146, 64, 14)
        impact_r = meta_p.add_run(f"Impact: {s.get('impact', 'Regional')}")
        impact_r.font.size = Pt(9)
        impact_r.font.color.rgb = RGBColor(22, 101, 52)
        
        # Why This Matters box
        why_p = doc.add_paragraph()
        why_label = why_p.add_run("WHY THIS MATTERS: ")
        why_label.font.size = Pt(9)
        why_label.font.bold = True
        why_label.font.color.rgb = RGBColor(180, 83, 9)
        why_text = why_p.add_run(s.get("summary", "Strategic development requiring executive attention."))
        why_text.font.size = Pt(10)
        
        # Link
        link_p = doc.add_paragraph()
        add_hyperlink(link_p, s["url"], "→ Read Full Article")
        link_p.space_after = Pt(16)
    
    doc.add_paragraph('─' * 70)
    
    # ─────────────────────────────────────────────────────────────
    # 04. COMPETITIVE & PEER SIGNALS
    # ─────────────────────────────────────────────────────────────
    
    add_section_header(doc, 4, "Competitive & Peer Signals")
    
    if competitive:
        for sig in competitive[:5]:
            cp = doc.add_paragraph(style='List Bullet')
            cr = cp.add_run(sig["headline"][:100])
            cr.font.size = Pt(10)
            src = cp.add_run(f"  ({sig['source']})")
            src.font.size = Pt(9)
            src.font.color.rgb = RGBColor(100, 116, 139)
    else:
        np = doc.add_paragraph()
        nr = np.add_run("No significant competitive signals in the last 24 hours.")
        nr.font.size = Pt(10)
        nr.font.italic = True
    
    doc.add_paragraph()
    
    # ─────────────────────────────────────────────────────────────
    # 05. TECHNOLOGY & AI RADAR
    # ─────────────────────────────────────────────────────────────
    
    add_section_header(doc, 5, "Technology & AI Radar")
    
    if tech_signals:
        for sig in tech_signals[:5]:
            tp = doc.add_paragraph(style='List Bullet')
            tr = tp.add_run(sig["headline"][:100])
            tr.font.size = Pt(10)
            if sig.get("implication"):
                imp = tp.add_run(f" — {sig['implication']}")
                imp.font.size = Pt(9)
                imp.font.color.rgb = RGBColor(124, 58, 237)
    else:
        np = doc.add_paragraph()
        nr = np.add_run("No significant technology signals in the last 24 hours.")
        nr.font.size = Pt(10)
        nr.font.italic = True
    
    doc.add_paragraph()
    
    # ─────────────────────────────────────────────────────────────
    # 06. RISK & THREAT WATCH
    # ─────────────────────────────────────────────────────────────
    
    add_section_header(doc, 6, "Risk & Threat Watch")
    
    risk_label = doc.add_paragraph()
    rl_run = risk_label.add_run("⚠ ALERTS FOR CRO/CISO ATTENTION")
    rl_run.font.size = Pt(9)
    rl_run.font.bold = True
    rl_run.font.color.rgb = RGBColor(220, 38, 38)
    
    if risks:
        for sig in risks[:5]:
            rp = doc.add_paragraph(style='List Bullet')
            rr = rp.add_run(sig["headline"])
            rr.font.size = Pt(10)
            if sig.get("severity") == "High":
                rr.font.bold = True
                rr.font.color.rgb = RGBColor(220, 38, 38)
            sev = rp.add_run(f"  [Severity: {sig.get('severity', 'Medium')}]")
            sev.font.size = Pt(9)
            sev.font.color.rgb = RGBColor(100, 116, 139)
    else:
        np = doc.add_paragraph()
        nr = np.add_run("No critical risk alerts in the last 24 hours.")
        nr.font.size = Pt(10)
        nr.font.italic = True
    
    doc.add_paragraph()
    
    # ─────────────────────────────────────────────────────────────
    # 07. CXO ACTION CUES
    # ─────────────────────────────────────────────────────────────
    
    doc.add_paragraph('─' * 70)
    add_section_header(doc, 7, "CXO Action Cues")
    
    # WATCH
    watch_p = doc.add_paragraph()
    watch_label = watch_p.add_run("👁 WATCH:  ")
    watch_label.font.size = Pt(10)
    watch_label.font.bold = True
    watch_label.font.color.rgb = RGBColor(251, 191, 36)
    watch_text = watch_p.add_run("Early indicators of regulatory shifts and competitive moves worth monitoring over the next 30 days")
    watch_text.font.size = Pt(10)
    
    # PREPARE
    prep_p = doc.add_paragraph()
    prep_label = prep_p.add_run("📋 PREPARE:  ")
    prep_label.font.size = Pt(10)
    prep_label.font.bold = True
    prep_label.font.color.rgb = RGBColor(251, 146, 60)
    prep_text = prep_p.add_run("Technology investments and capability builds that peers are executing; assess relevance to your strategy")
    prep_text.font.size = Pt(10)
    
    # ACT
    act_p = doc.add_paragraph()
    act_label = act_p.add_run("⚡ ACT:  ")
    act_label.font.size = Pt(10)
    act_label.font.bold = True
    act_label.font.color.rgb = RGBColor(248, 113, 113)
    act_text = act_p.add_run("Immediate review of cyber risk posture and compliance readiness based on today's threat signals")
    act_text.font.size = Pt(10)
    
    # ─────────────────────────────────────────────────────────────
    # FOOTER
    # ─────────────────────────────────────────────────────────────
    
    doc.add_paragraph()
    doc.add_paragraph('─' * 70)
    
    footer_brand = doc.add_paragraph()
    footer_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fb_run = footer_brand.add_run('CONTINUUM INTELLIGENCE')
    fb_run.font.size = Pt(10)
    fb_run.font.color.rgb = RGBColor(59, 130, 246)
    fb_run.font.bold = True
    
    sources_p = doc.add_paragraph()
    sources_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src_run = sources_p.add_run(
        'Sources: Insurance Journal · Digital Insurance · Dark Reading · Reuters · '
        'BBC Business · CNBC · Financial Times · TechCrunch'
    )
    src_run.font.size = Pt(9)
    src_run.font.color.rgb = RGBColor(120, 120, 120)
    
    disc_p = doc.add_paragraph()
    disc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_run = disc_p.add_run(
        'This intelligence brief is prepared for executive decision-makers. '
        'Information derived from sources believed reliable but not guaranteed. '
        '© 2025 Continuum. All rights reserved.'
    )
    disc_run.font.size = Pt(8)
    disc_run.font.color.rgb = RGBColor(150, 150, 150)
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"CXO_Intel_Brief_{now.strftime('%Y%m%d')}.docx"
    return Response(
        buffer.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


def generate_executive_bullets(stories, themes, risks, tech_signals):
    """Generate executive snapshot bullets."""
    bullets = []
    
    # Key signal from themes
    if themes:
        bullets.append({
            "text": f"Key Signal: {themes[0]['name']} dominating today's intelligence with {len(stories)} high-priority items",
            "cxo": "CEO"
        })
    
    # Risk alert
    if risks:
        bullets.append({
            "text": f"Risk Alert: {len(risks)} cybersecurity/risk developments requiring board-level visibility",
            "cxo": "CRO/CISO"
        })
    
    # Tech radar
    if tech_signals:
        bullets.append({
            "text": f"Tech Radar: {len(tech_signals)} AI/technology developments with strategic implications",
            "cxo": "CIO/CDO"
        })
    
    # Regulatory signal from stories
    reg_stories = [s for s in stories if "Regulation" in s.get("theme", "")]
    if reg_stories:
        bullets.append({
            "text": f"Compliance Watch: {len(reg_stories)} regulatory developments requiring legal/compliance review",
            "cxo": "CLO/CCO"
        })
    
    if not bullets:
        bullets.append({
            "text": "Market Scan: Limited high-priority signals in the last 24 hours; routine monitoring continues",
            "cxo": "All CXOs"
        })
    
    return bullets[:5]
